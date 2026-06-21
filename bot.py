import os
import logging
import asyncio
import subprocess
import tempfile
import shutil
import json
from pathlib import Path
from typing import Optional

from PIL import Image
import img2pdf
from pypdf import PdfReader, PdfWriter
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from telegram.constants import ParseMode

logging.basicConfig(format=”%(asctime)s | %(levelname)s | %(name)s | %(message)s”, level=logging.INFO)
logger = logging.getLogger(**name**)

BOT_TOKEN = os.environ.get(“BOT_TOKEN”, “YOUR_BOT_TOKEN_HERE”)
TEMP_DIR = Path(tempfile.gettempdir()) / “fileflux”
TEMP_DIR.mkdir(exist_ok=True)

IMAGE_FORMATS = [“JPEG”, “PNG”, “WEBP”, “BMP”, “GIF”, “TIFF”, “ICO”, “PDF”]
VIDEO_FORMATS = [“MP4”, “AVI”, “MKV”, “MOV”, “WEBM”, “GIF”, “MP3”, “AAC”, “WAV”, “VIDNOTE”]
AUDIO_FORMATS = [“MP3”, “WAV”, “AAC”, “OGG”, “FLAC”, “M4A”, “OPUS”]
DOC_FORMATS = [“PDF”, “TXT”, “DOCX”]

def detect_type(mime, filename):
if mime:
if mime.startswith(“image/”): return “image”
if mime.startswith(“video/”): return “video”
if mime.startswith(“audio/”): return “audio”
if mime in (“application/pdf”, “text/plain”,
“application/vnd.openxmlformats-officedocument.wordprocessingml.document”,
“application/msword”):
return “doc”
if filename:
ext = Path(filename).suffix.lower()
if ext in {”.jpg”, “.jpeg”, “.png”, “.webp”, “.bmp”, “.gif”, “.tiff”, “.tif”, “.ico”}: return “image”
if ext in {”.mp4”, “.avi”, “.mkv”, “.mov”, “.webm”, “.flv”, “.wmv”, “.3gp”, “.m4v”}: return “video”
if ext in {”.mp3”, “.wav”, “.aac”, “.ogg”, “.flac”, “.m4a”, “.opus”, “.wma”}: return “audio”
if ext in {”.pdf”, “.txt”, “.docx”, “.doc”}: return “doc”
return “unknown”

def user_dir(user_id):
d = TEMP_DIR / str(user_id)
d.mkdir(exist_ok=True)
return d

def cleanup(user_id):
d = user_dir(user_id)
shutil.rmtree(d, ignore_errors=True)
d.mkdir(exist_ok=True)

def format_buttons(file_type):
formats_map = {
“image”: IMAGE_FORMATS,
“video”: VIDEO_FORMATS,
“audio”: AUDIO_FORMATS,
“doc”: DOC_FORMATS
}
formats = formats_map.get(file_type, [])
buttons = []
row = []
for fmt in formats:
row.append(InlineKeyboardButton(fmt, callback_data=“convert:” + fmt))
if len(row) == 3:
buttons.append(row)
row = []
if row:
buttons.append(row)
if file_type == “image”:
buttons.append([
InlineKeyboardButton(“📐 Resize”, callback_data=“tool:resize”),
InlineKeyboardButton(“🗜 Compress”, callback_data=“tool:compress”),
])
if file_type == “video”:
buttons.append([
InlineKeyboardButton(“✂️ Trim”, callback_data=“tool:trim”),
InlineKeyboardButton(“🔇 Mute”, callback_data=“tool:mute”),
InlineKeyboardButton(“🖼 Thumbnail”, callback_data=“tool:thumbnail”),
])
buttons.append([
InlineKeyboardButton(“🗜 Compress”, callback_data=“tool:compress”),
InlineKeyboardButton(“📊 File Info”, callback_data=“tool:info”),
])
if file_type == “audio”:
buttons.append([InlineKeyboardButton(“📊 File Info”, callback_data=“tool:info”)])
if file_type == “doc”:
buttons.append([InlineKeyboardButton(“📝 Merge PDFs”, callback_data=“tool:merge_pdf”)])
buttons.append([InlineKeyboardButton(“❌ Cancel”, callback_data=“cancel”)])
return InlineKeyboardMarkup(buttons)

def get_file_info(src):
cmd = [“ffprobe”, “-v”, “quiet”, “-print_format”, “json”, “-show_format”, “-show_streams”, str(src)]
result = subprocess.run(cmd, capture_output=True, text=True)
try:
data = json.loads(result.stdout)
fmt = data.get(“format”, {})
streams = data.get(“streams”, [])
video_stream = next((s for s in streams if s.get(“codec_type”) == “video”), None)
audio_stream = next((s for s in streams if s.get(“codec_type”) == “audio”), None)
info = {
“size”: int(fmt.get(“size”, 0)),
“duration”: float(fmt.get(“duration”, 0)),
“format”: fmt.get(“format_long_name”, “Unknown”),
}
if video_stream:
info[“width”] = video_stream.get(“width”, 0)
info[“height”] = video_stream.get(“height”, 0)
info[“video_codec”] = video_stream.get(“codec_name”, “Unknown”)
try:
info[“fps”] = eval(video_stream.get(“r_frame_rate”, “0/1”))
except Exception:
info[“fps”] = 0
if audio_stream:
info[“audio_codec”] = audio_stream.get(“codec_name”, “Unknown”)
info[“sample_rate”] = audio_stream.get(“sample_rate”, “Unknown”)
return info
except Exception:
return {}

def convert_image(src, target_fmt, out_dir):
fmt = target_fmt.upper()
if fmt == “PDF”:
out = out_dir / (src.stem + “.pdf”)
with open(out, “wb”) as f:
f.write(img2pdf.convert(str(src)))
return out
pil_fmt = “JPEG” if fmt == “JPG” else fmt
ext_map = {“JPEG”: “jpg”, “PNG”: “png”, “WEBP”: “webp”, “BMP”: “bmp”, “GIF”: “gif”, “TIFF”: “tiff”, “ICO”: “ico”}
ext = ext_map.get(pil_fmt, pil_fmt.lower())
out = out_dir / (src.stem + “.” + ext)
img = Image.open(src)
if pil_fmt in (“JPEG”, “BMP”, “ICO”) and img.mode in (“RGBA”, “P”, “LA”):
img = img.convert(“RGB”)
save_kwargs = {“quality”: 92, “optimize”: True} if pil_fmt == “JPEG” else {}
img.save(out, pil_fmt, **save_kwargs)
return out

def resize_image(src, width, height, out_dir):
out = out_dir / (src.stem + “_resized” + src.suffix)
img = Image.open(src)
if width == 0:
width = int(img.width * height / img.height)
elif height == 0:
height = int(img.height * width / img.width)
img = img.resize((width, height), Image.LANCZOS)
img.save(out)
return out

def compress_image(src, quality, out_dir):
out = out_dir / (src.stem + “_compressed.jpg”)
img = Image.open(src)
if img.mode in (“RGBA”, “P”, “LA”):
img = img.convert(“RGB”)
img.save(out, “JPEG”, quality=quality, optimize=True)
return out

def convert_video(src, target_fmt, out_dir):
fmt = target_fmt.upper()
if fmt == “GIF”:
out = out_dir / (src.stem + “.gif”)
vf = “fps=10,scale=480:-1:flags=lanczos,split[s0][s1];[s0]palettegen[p];[s1][p]paletteuse”
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src), “-vf”, vf, “-loop”, “0”, str(out)], check=True, capture_output=True)
return out
if fmt == “VIDNOTE”:
out = out_dir / (src.stem + “_note.mp4”)
vf = “crop=min(iw\,ih):min(iw\,ih),scale=384:384”
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src), “-t”, “60”, “-vf”, vf, “-c:v”, “libx264”, “-preset”, “fast”, “-crf”, “28”, “-c:a”, “aac”, “-b:a”, “64k”, str(out)], check=True, capture_output=True)
return out
if fmt in (“MP3”, “AAC”, “WAV”, “OGG”, “FLAC”, “M4A”, “OPUS”):
return convert_audio(src, fmt, out_dir)
ext_map = {“MP4”: “mp4”, “AVI”: “avi”, “MKV”: “mkv”, “MOV”: “mov”, “WEBM”: “webm”}
ext = ext_map.get(fmt, fmt.lower())
out = out_dir / (src.stem + “.” + ext)
codec_map = {
“MP4”: [”-c:v”, “libx264”, “-preset”, “fast”, “-crf”, “23”, “-c:a”, “aac”],
“AVI”: [”-c:v”, “libxvid”, “-qscale:v”, “3”, “-c:a”, “libmp3lame”],
“MKV”: [”-c:v”, “libx264”, “-preset”, “fast”, “-crf”, “23”, “-c:a”, “aac”],
“MOV”: [”-c:v”, “libx264”, “-preset”, “fast”, “-crf”, “23”, “-c:a”, “aac”],
“WEBM”: [”-c:v”, “libvpx-vp9”, “-crf”, “30”, “-b:v”, “0”, “-c:a”, “libopus”],
}
codecs = codec_map.get(fmt, [”-c”, “copy”])
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src)] + codecs + [str(out)], check=True, capture_output=True)
return out

def trim_video(src, start, end, out_dir):
out = out_dir / (src.stem + “_trimmed.mp4”)
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src), “-ss”, start, “-to”, end, “-c:v”, “libx264”, “-preset”, “fast”, “-crf”, “23”, “-c:a”, “aac”, str(out)], check=True, capture_output=True)
return out

def mute_video(src, out_dir):
out = out_dir / (src.stem + “_muted.mp4”)
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src), “-c:v”, “copy”, “-an”, str(out)], check=True, capture_output=True)
return out

def extract_thumbnail(src, timestamp, out_dir):
out = out_dir / (src.stem + “_thumb.jpg”)
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src), “-ss”, timestamp, “-vframes”, “1”, “-q:v”, “2”, str(out)], check=True, capture_output=True)
return out

def compress_video(src, out_dir):
out = out_dir / (src.stem + “_compressed.mp4”)
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src), “-c:v”, “libx264”, “-preset”, “medium”, “-crf”, “28”, “-c:a”, “aac”, “-b:a”, “128k”, str(out)], check=True, capture_output=True)
return out

def convert_audio(src, target_fmt, out_dir):
fmt = target_fmt.upper()
ext_map = {“MP3”: “mp3”, “WAV”: “wav”, “AAC”: “aac”, “OGG”: “ogg”, “FLAC”: “flac”, “M4A”: “m4a”, “OPUS”: “opus”}
ext = ext_map.get(fmt, fmt.lower())
out = out_dir / (src.stem + “.” + ext)
codec_map = {
“MP3”: [”-c:a”, “libmp3lame”, “-q:a”, “2”],
“WAV”: [”-c:a”, “pcm_s16le”],
“AAC”: [”-c:a”, “aac”, “-b:a”, “192k”],
“OGG”: [”-c:a”, “libvorbis”, “-q:a”, “4”],
“FLAC”: [”-c:a”, “flac”],
“M4A”: [”-c:a”, “aac”, “-b:a”, “192k”],
“OPUS”: [”-c:a”, “libopus”, “-b:a”, “128k”],
}
codecs = codec_map.get(fmt, [”-c:a”, “copy”])
subprocess.run([“ffmpeg”, “-y”, “-i”, str(src)] + codecs + [str(out)], check=True, capture_output=True)
return out

def convert_doc(src, target_fmt, out_dir):
fmt = target_fmt.upper()
src_ext = src.suffix.lower()
if src_ext == “.pdf” and fmt == “TXT”:
from pdfminer.high_level import extract_text
out = out_dir / (src.stem + “.txt”)
out.write_text(extract_text(str(src)), encoding=“utf-8”)
return out
if src_ext == “.pdf” and fmt == “DOCX”:
from pdfminer.high_level import extract_text
out = out_dir / (src.stem + “.docx”)
doc = Document()
for line in extract_text(str(src)).split(”\n”):
doc.add_paragraph(line)
doc.save(str(out))
return out
if src_ext == “.txt” and fmt == “PDF”:
out = out_dir / (src.stem + “.pdf”)
content = src.read_text(encoding=“utf-8”, errors=“replace”)
pdf_doc = SimpleDocTemplate(str(out), pagesize=A4)
styles = getSampleStyleSheet()
story = []
for line in content.split(”\n”):
story.append(Paragraph(line or “ “, styles[“Normal”]))
story.append(Spacer(1, 4))
pdf_doc.build(story)
return out
if src_ext == “.txt” and fmt == “DOCX”:
out = out_dir / (src.stem + “.docx”)
doc = Document()
for line in src.read_text(encoding=“utf-8”, errors=“replace”).split(”\n”):
doc.add_paragraph(line)
doc.save(str(out))
return out
if src_ext == “.docx” and fmt == “PDF”:
out = out_dir / (src.stem + “.pdf”)
doc_in = Document(str(src))
pdf_doc = SimpleDocTemplate(str(out), pagesize=A4)
styles = getSampleStyleSheet()
story = []
for line in “\n”.join(p.text for p in doc_in.paragraphs).split(”\n”):
story.append(Paragraph(line or “ “, styles[“Normal”]))
story.append(Spacer(1, 4))
pdf_doc.build(story)
return out
if src_ext == “.docx” and fmt == “TXT”:
out = out_dir / (src.stem + “.txt”)
doc_in = Document(str(src))
out.write_text(”\n”.join(p.text for p in doc_in.paragraphs), encoding=“utf-8”)
return out
raise ValueError(“Conversion “ + src_ext.upper() + “ to “ + fmt + “ not supported”)

def merge_pdfs(paths, out_dir):
out = out_dir / “merged.pdf”
writer = PdfWriter()
for p in paths:
reader = PdfReader(str(p))
for page in reader.pages:
writer.add_page(page)
with open(out, “wb”) as f:
writer.write(f)
return out

async def cmd_start(update, ctx):
await update.message.reply_text(
“🤖 *Welcome to FileFlux v2!*\n\n”
“Send me any file and I will convert or process it.\n\n”
“🖼 Images: JPEG PNG WEBP BMP GIF TIFF ICO PDF + Resize + Compress\n”
“🎬 Videos: MP4 AVI MKV MOV WEBM GIF + Trim Mute Thumbnail Compress\n”
“🎵 Audio: MP3 WAV AAC OGG FLAC M4A OPUS\n”
“📄 Docs: PDF TXT DOCX + PDF Merge\n\n”
“Just send a file to get started.\n”
“/help for more info”,
parse_mode=ParseMode.MARKDOWN
)

async def cmd_help(update, ctx):
await update.message.reply_text(
“ℹ️ *FileFlux v2 Help*\n\n”
“1. Send any file\n”
“2. Pick a format or tool\n”
“3. Follow any prompts\n”
“4. Get your result\n\n”
“Tools available:\n”
“✂️ Trim - cut video, send start then end time e.g. 0:10\n”
“📐 Resize - send dimensions e.g. 1280x720\n”
“🗜 Compress - reduce file size\n”
“🔇 Mute - remove audio from video\n”
“🖼 Thumbnail - extract frame at a timestamp\n”
“📝 Merge PDFs - send PDFs one by one then merge\n”
“📊 File Info - see codec resolution duration size\n\n”
“/cancel - cancel current operation”,
parse_mode=ParseMode.MARKDOWN
)

async def cmd_cancel(update, ctx):
cleanup(update.effective_user.id)
ctx.user_data.clear()
await update.message.reply_text(“Cancelled. Send a new file whenever you are ready.”)

async def cmd_formats(update, ctx):
await update.message.reply_text(
“⭐ *Supported Formats*\n\n”
“🖼 Image input: JPG PNG WEBP BMP GIF TIFF ICO\n”
“   Output: JPEG PNG WEBP BMP GIF TIFF ICO PDF\n\n”
“🎬 Video input: MP4 AVI MKV MOV WEBM FLV 3GP\n”
“   Output: MP4 AVI MKV MOV WEBM GIF MP3 AAC WAV VIDNOTE\n\n”
“🎵 Audio input: MP3 WAV AAC OGG FLAC M4A OPUS WMA\n”
“   Output: MP3 WAV AAC OGG FLAC M4A OPUS\n\n”
“📄 Doc input: PDF TXT DOCX\n”
“   Output: PDF TXT DOCX”,
parse_mode=ParseMode.MARKDOWN
)

async def handle_file(update, ctx):
msg = update.message
uid = update.effective_user.id

```
if ctx.user_data.get("tool") == "merge_pdf":
    if msg.document and msg.document.mime_type == "application/pdf":
        merge_files = ctx.user_data.get("merge_files", [])
        tg_file = await ctx.bot.get_file(msg.document.file_id)
        dest = user_dir(uid) / ("merge_" + str(len(merge_files)) + ".pdf")
        await tg_file.download_to_drive(str(dest))
        merge_files.append(str(dest))
        ctx.user_data["merge_files"] = merge_files
        count = len(merge_files)
        keyboard = InlineKeyboardMarkup([[
            InlineKeyboardButton("📝 Merge " + str(count) + " PDFs now", callback_data="do_merge"),
            InlineKeyboardButton("❌ Cancel", callback_data="cancel"),
        ]])
        await msg.reply_text("Got PDF " + str(count) + ". Send more or tap merge.", reply_markup=keyboard)
        return

file_obj = None
mime = None
filename = None

if msg.photo:
    file_obj = msg.photo[-1]
    mime = "image/jpeg"
    filename = "photo.jpg"
elif msg.video:
    file_obj = msg.video
    mime = msg.video.mime_type
    filename = msg.video.file_name or "video.mp4"
elif msg.audio:
    file_obj = msg.audio
    mime = msg.audio.mime_type
    filename = msg.audio.file_name or "audio.mp3"
elif msg.voice:
    file_obj = msg.voice
    mime = msg.voice.mime_type or "audio/ogg"
    filename = "voice.ogg"
elif msg.video_note:
    file_obj = msg.video_note
    mime = "video/mp4"
    filename = "video_note.mp4"
elif msg.document:
    file_obj = msg.document
    mime = msg.document.mime_type
    filename = msg.document.file_name or "file"
elif msg.animation:
    file_obj = msg.animation
    mime = "video/mp4"
    filename = msg.animation.file_name or "animation.mp4"

if file_obj is None:
    await msg.reply_text("❌ Unsupported file.")
    return

file_size = getattr(file_obj, "file_size", None)
if file_size and file_size > 50 * 1024 * 1024:
    await msg.reply_text("❌ File too large. Max 50 MB.")
    return

ftype = detect_type(mime, filename)
if ftype == "unknown":
    await msg.reply_text("❌ Unsupported file type.")
    return

ctx.user_data["file_id"] = file_obj.file_id
ctx.user_data["filename"] = filename
ctx.user_data["file_type"] = ftype
ctx.user_data["tool"] = None

labels = {"image": "🖼 Image", "video": "🎬 Video", "audio": "🎵 Audio", "doc": "📄 Document"}
label = labels.get(ftype, "File")
size_str = " (" + str(file_size // 1024) + " KB)" if file_size else ""

await msg.reply_text(
    "⏳ *" + label + " received*" + size_str + "\n" + "`" + filename + "`" + "\n\nChoose what to do:",
    parse_mode=ParseMode.MARKDOWN,
    reply_markup=format_buttons(ftype),
)
```

async def handle_text(update, ctx):
uid = update.effective_user.id
text = update.message.text.strip()
tool = ctx.user_data.get(“tool”)

```
if tool == "trim_waiting_start":
    ctx.user_data["trim_start"] = text
    ctx.user_data["tool"] = "trim_waiting_end"
    await update.message.reply_text("Now send the end time (e.g. 1:30):")
    return

if tool == "trim_waiting_end":
    ctx.user_data["trim_end"] = text
    await update.message.reply_text("⏳ Trimming video...")
    await run_tool(update.message, ctx, uid, "trim")
    return

if tool == "resize_waiting":
    try:
        parts = text.lower().replace("x", " ").split()
        w, h = int(parts[0]), int(parts[1])
        ctx.user_data["resize_w"] = w
        ctx.user_data["resize_h"] = h
        await update.message.reply_text("⏳ Resizing...")
        await run_tool(update.message, ctx, uid, "resize")
    except Exception:
        await update.message.reply_text("❌ Send dimensions like 1280x720")
    return

if tool == "thumbnail_waiting":
    ctx.user_data["thumb_time"] = text
    await update.message.reply_text("⏳ Extracting thumbnail...")
    await run_tool(update.message, ctx, uid, "thumbnail")
    return

if tool == "compress_waiting":
    try:
        quality = int(text)
        if not 1 <= quality <= 100:
            raise ValueError
        ctx.user_data["compress_quality"] = quality
        await update.message.reply_text("⏳ Compressing...")
        await run_tool(update.message, ctx, uid, "compress")
    except Exception:
        await update.message.reply_text("❌ Send a number between 1 and 100.")
    return
```

async def run_tool(msg, ctx, uid, tool_name):
file_id = ctx.user_data.get(“file_id”)
filename = ctx.user_data.get(“filename”, “file”)
file_type = ctx.user_data.get(“file_type”)
work_dir = user_dir(uid)

```
try:
    tg_file = await msg.get_bot().get_file(file_id)
    src_path = work_dir / filename
    await tg_file.download_to_drive(str(src_path))
except Exception as e:
    await msg.reply_text("❌ Download failed: " + str(e)[:100])
    return

try:
    out_path = None

    if tool_name == "trim":
        start = ctx.user_data.get("trim_start", "0:00")
        end = ctx.user_data.get("trim_end", "0:30")
        out_path = await asyncio.get_event_loop().run_in_executor(None, trim_video, src_path, start, end, work_dir)

    elif tool_name == "mute":
        out_path = await asyncio.get_event_loop().run_in_executor(None, mute_video, src_path, work_dir)

    elif tool_name == "thumbnail":
        ts = ctx.user_data.get("thumb_time", "0:00")
        out_path = await asyncio.get_event_loop().run_in_executor(None, extract_thumbnail, src_path, ts, work_dir)

    elif tool_name == "resize":
        w = ctx.user_data.get("resize_w", 0)
        h = ctx.user_data.get("resize_h", 0)
        out_path = await asyncio.get_event_loop().run_in_executor(None, resize_image, src_path, w, h, work_dir)

    elif tool_name == "compress":
        if file_type == "image":
            quality = ctx.user_data.get("compress_quality", 60)
            out_path = await asyncio.get_event_loop().run_in_executor(None, compress_image, src_path, quality, work_dir)
        else:
            out_path = await asyncio.get_event_loop().run_in_executor(None, compress_video, src_path, work_dir)

    elif tool_name == "info":
        info = await asyncio.get_event_loop().run_in_executor(None, get_file_info, src_path)
        size_mb = info.get("size", 0) / 1024 / 1024
        dur = info.get("duration", 0)
        minutes = int(dur // 60)
        seconds = int(dur % 60)
        lines = [
            "📊 *File Info*",
            "📁 Size: `" + str(round(size_mb, 2)) + " MB`",
            "⏱ Duration: `" + str(minutes) + ":" + str(seconds).zfill(2) + "`",
            "🎞 Format: `" + info.get("format", "Unknown") + "`",
        ]
        if "width" in info:
            lines.append("📐 Resolution: `" + str(info["width"]) + "x" + str(info["height"]) + "`")
            lines.append("🎬 Video codec: `" + info.get("video_codec", "N/A") + "`")
            lines.append("🎯 FPS: `" + str(round(info.get("fps", 0), 1)) + "`")
        if "audio_codec" in info:
            lines.append("🎵 Audio codec: `" + info.get("audio_codec", "N/A") + "`")
            lines.append("🔊 Sample rate: `" + str(info.get("sample_rate", "N/A")) + " Hz`")
        await msg.reply_text("\n".join(lines), parse_mode=ParseMode.MARKDOWN)
        cleanup(uid)
        ctx.user_data.clear()
        return

    if out_path:
        out_size = out_path.stat().st_size
        caption = "✅ Done! `" + out_path.name + "` (" + str(out_size // 1024) + " KB)"
        ext = out_path.suffix.lower()
        with open(out_path, "rb") as f:
            if tool_name == "thumbnail" or (file_type == "image" and ext != ".pdf"):
                await msg.reply_photo(photo=f, caption=caption, parse_mode=ParseMode.MARKDOWN)
            elif file_type == "video" and ext in (".mp4", ".mov", ".webm"):
                await msg.reply_video(video=f, caption=caption, parse_mode=ParseMode.MARKDOWN)
            else:
                await msg.reply_document(document=f, caption=caption, parse_mode=ParseMode.MARKDOWN)

except Exception as e:
    logger.error("Tool error: " + str(e))
    await msg.reply_text("❌ Error: " + str(e)[:200])
finally:
    cleanup(uid)
    ctx.user_data.clear()
```

async def handle_callback(update, ctx):
query = update.callback_query
await query.answer()
uid = query.from_user.id

```
if query.data == "cancel":
    cleanup(uid)
    ctx.user_data.clear()
    await query.edit_message_text("Cancelled.")
    return

if query.data == "do_merge":
    merge_files = ctx.user_data.get("merge_files", [])
    if len(merge_files) < 2:
        await query.edit_message_text("❌ Send at least 2 PDFs to merge.")
        return
    await query.edit_message_text("⏳ Merging " + str(len(merge_files)) + " PDFs...")
    try:
        out_path = await asyncio.get_event_loop().run_in_executor(None, merge_pdfs, [Path(p) for p in merge_files], user_dir(uid))
        out_size = out_path.stat().st_size
        with open(out_path, "rb") as f:
            await query.message.reply_document(document=f, caption="✅ Merged! (" + str(out_size // 1024) + " KB)")
        await query.edit_message_text("✅ Merge complete!")
    except Exception as e:
        await query.edit_message_text("❌ Merge failed: " + str(e)[:200])
    finally:
        cleanup(uid)
        ctx.user_data.clear()
    return

if query.data.startswith("tool:"):
    tool = query.data.split(":", 1)[1]
    file_id = ctx.user_data.get("file_id")
    file_type = ctx.user_data.get("file_type")

    if not file_id:
        await query.edit_message_text("❌ Session expired. Send your file again.")
        return

    if tool == "trim":
        ctx.user_data["tool"] = "trim_waiting_start"
        await query.edit_message_text("✂️ Send the start time (e.g. 0:10 or 1:30):")
        return

    if tool == "resize":
        ctx.user_data["tool"] = "resize_waiting"
        await query.edit_message_text("📐 Send dimensions as widthxheight e.g. 1280x720 or 800x0 for auto height:")
        return

    if tool == "thumbnail":
        ctx.user_data["tool"] = "thumbnail_waiting"
        await query.edit_message_text("🖼 Send the timestamp to capture (e.g. 0:05 or 1:23):")
        return

    if tool == "compress":
        if file_type == "image":
            ctx.user_data["tool"] = "compress_waiting"
            await query.edit_message_text("🗜 Send quality level 1-100. 80 = high quality, 60 = balanced, 40 = smaller size:")
            return
        else:
            ctx.user_data["tool"] = "compress"
            await query.edit_message_text("⏳ Compressing video...")
            await run_tool(query.message, ctx, uid, "compress")
            return

    if tool == "mute":
        ctx.user_data["tool"] = "mute"
        await query.edit_message_text("⏳ Removing audio...")
        await run_tool(query.message, ctx, uid, "mute")
        return

    if tool == "info":
        ctx.user_data["tool"] = "info"
        await query.edit_message_text("⏳ Reading file info...")
        await run_tool(query.message, ctx, uid, "info")
        return

    if tool == "merge_pdf":
        ctx.user_data["tool"] = "merge_pdf"
        ctx.user_data["merge_files"] = []
        await query.edit_message_text("📝 PDF Merge Mode. Send PDFs one by one, then tap merge when ready.")
        return

if query.data.startswith("convert:"):
    target_fmt = query.data.split(":", 1)[1]
    file_id = ctx.user_data.get("file_id")
    filename = ctx.user_data.get("filename", "file")
    file_type = ctx.user_data.get("file_type")

    if not file_id:
        await query.edit_message_text("❌ Session expired. Send your file again.")
        return

    await query.edit_message_text("⏳ Converting to " + target_fmt + "...")

    work_dir = user_dir(uid)
    try:
        tg_file = await ctx.bot.get_file(file_id)
        src_path = work_dir / filename
        await tg_file.download_to_drive(str(src_path))
    except Exception:
        await query.edit_message_text("❌ Download failed.")
        return

    try:
        converters = {"image": convert_image, "video": convert_video, "audio": convert_audio, "doc": convert_doc}
        converter = converters.get(file_type)
        if not converter:
            raise ValueError("No converter for: " + str(file_type))
        out_path = await asyncio.get_event_loop().run_in_executor(None, converter, src_path, target_fmt, work_dir)
    except subprocess.CalledProcessError:
        await query.edit_message_text("❌ Conversion failed. File may be corrupted or format unsupported.")
        cleanup(uid)
        return
    except Exception as e:
        await query.edit_message_text("❌ Error: " + str(e)[:200])
        cleanup(uid)
        return

    try:
        out_size = out_path.stat().st_size
        caption = "✅ Done! `" + out_path.name + "` (" + str(out_size // 1024) + " KB)"
        with open(out_path, "rb") as f:
            ext = out_path.suffix.lower()
            if "note" in out_path.stem:
                await query.message.reply_video_note(video_note=f)
            elif file_type == "image" and ext != ".pdf":
                await query.message.reply_photo(photo=f, caption=caption, parse_mode=ParseMode.MARKDOWN)
            elif file_type == "video" and ext in (".mp4", ".mov", ".webm"):
                await query.message.reply_video(video=f, caption=caption, parse_mode=ParseMode.MARKDOWN)
            elif file_type == "audio" and ext in (".mp3", ".ogg", ".m4a", ".aac"):
                await query.message.reply_audio(audio=f, caption=caption, parse_mode=ParseMode.MARKDOWN)
            else:
                await query.message.reply_document(document=f, caption=caption, parse_mode=ParseMode.MARKDOWN)
        await query.edit_message_text("✅ Done! Your " + target_fmt + " file is ready.")
    except Exception as e:
        await query.edit_message_text("❌ Converted but failed to send: " + str(e)[:200])
    finally:
        cleanup(uid)
        ctx.user_data.clear()
```

def main():
if BOT_TOKEN == “YOUR_BOT_TOKEN_HERE”:
print(“Set your BOT_TOKEN environment variable first!”)
return
app = Application.builder().token(BOT_TOKEN).build()
app.add_handler(CommandHandler(“start”, cmd_start))
app.add_handler(CommandHandler(“help”, cmd_help))
app.add_handler(CommandHandler(“formats”, cmd_formats))
app.add_handler(CommandHandler(“cancel”, cmd_cancel))
app.add_handler(MessageHandler(
filters.PHOTO | filters.VIDEO | filters.AUDIO | filters.VOICE |
filters.VIDEO_NOTE | filters.Document.ALL | filters.ANIMATION,
handle_file
))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
app.add_handler(CallbackQueryHandler(handle_callback))
print(“FileFlux v2 is running…”)
app.run_polling(drop_pending_updates=True)

if **name** == “**main**”:
main()
